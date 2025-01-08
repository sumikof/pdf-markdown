import os

from spire.pdf import PdfDocument
from spire.pdf import FileFormat

pdf = PdfDocument()

dirpath = os.path.abspath(os.path.dirname(__file__))
pdf.LoadFromFile(os.path.join(dirpath,"140120250107547220.pdf"))

pdf.SaveToFile("spi1.doc", FileFormat.DOC)

pdf.SaveToFile("spi2.docx", FileFormat.DOCX)

pdf.Close()

# PDFをWordに変換し、HTML形式に変換
def convert_word_to_html(word_file, output_html):
    from docx import Document
    doc = Document(word_file)
    html_content = ""
    for paragraph in doc.paragraphs:
        html_content += f"<p>{paragraph.text}</p>\n"
    
    # HTMLファイルに保存
    with open(output_html, "w", encoding="utf-8") as html_file:
        html_file.write(html_content)

convert_word_to_html("spi2.docx","spi2.html")